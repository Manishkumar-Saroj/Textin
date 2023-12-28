import streamlit as st
import os
import requests
from pydub import AudioSegment
import io
from time import sleep
import sqlite3
import hashlib
from io import BytesIO
from docx import Document
from docx.shared import Inches


#--------------------------------SET PAGE CONFIGRATION---------------------------- 
st.set_page_config(
    page_title="Textin",
    layout="wide",    
    initial_sidebar_state="auto",
    
)

# Container for the top section to insert logo and basic detail about website
top_container = st.container()

# Two columns within the container
col1, col2 = top_container.columns([3, 1])


col1.header('Hello, 👋')
col1.subheader('Welcome to Textin')
col1.write("Welcome to our transcribing website! Simply upload your audio or video file, and we'll transcribe it into text for you. Our service is quick, precise, and free to use. Try it out today!")
# Image is for logo of the website.
col2.image('images\logo.png', use_column_width=True) 



   



#--------------------------------------CODE FOR TO HIDE MAINMENU AND FOOTER----------------------------
hide_streamlit_style = """
<style>
#MainMenu {visibility: hidden;}
footer {visibility: hidden;}
.streamlit-header {background-color: transparent;}
</style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)




#-------------------------------------SQLITE CONNECTION---------------------------------------------------
# Initialize SQLite connection
conn = sqlite3.connect('database.db')
c = conn.cursor()

# Create users table if not exists
c.execute('''CREATE TABLE IF NOT EXISTS users
             (username VARCHAR(10) PRIMARY KEY, password_hash VARCHAR(10), mobile_number VARCHAR(20))''')

# Create transcripts table if not exists
c.execute('''CREATE TABLE IF NOT EXISTS transcripts
             (id INTEGER PRIMARY KEY AUTOINCREMENT,
             username VARCHAR(10),
             filename VARCHAR(255),
             language VARCHAR(10),                          
             text TEXT,             
             transcript_status VARCHAR(20))''')

#-------------------------------------------------------------------------------------------------------------
# Function to create password hash using sha256
def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

# Function to check if password hash matches
def verify_password(password, hash):
    return hash_password(password) == hash

# Function to check if mobile number exists in database
def check_mobile_number(mobile_number):
    c.execute("SELECT * FROM users WHERE mobile_number=?", (mobile_number,))
    user = c.fetchone()
    return user[0] if user else None


#----------------------------------------------------------------------------------------------------------
# Login and signup forms
page = st.sidebar.selectbox('Select page', ['Login', 'Signup', 'Reset Password/Username'])
st.sidebar.write(page)


#----------------------------------------LOGIN FORM--------------------------------------------------
if page == 'Login':
    username = st.sidebar.text_input('Username')
    password = st.sidebar.text_input('Password', type='password')
    if st.sidebar.button('Login'):
        if not username or not password:
            st.sidebar.warning('Please enter both username and password.')
        else:
            # Check if user exists in database
            c.execute("SELECT * FROM users WHERE username=?", (username,))
            user = c.fetchone()
            if user and verify_password(password, user[1]):
                # Show main app if login is successful
                st.sidebar.success('Logged in successfully!')
                app_state = st.experimental_get_query_params()
                app_state['username'] = username
                st.experimental_set_query_params(**app_state)
            else:
                st.sidebar.error('Incorrect username or password')


#----------------------------------------SIGNUP FORM---------------------------------------------------            
elif page == 'Signup':
    username = st.sidebar.text_input('Username')
    password = st.sidebar.text_input('Password', type='password')
    password_confirm = st.sidebar.text_input('Confirm Password', type='password')
    mobile_number = st.sidebar.text_input('Mobile Number')

    # Check if username and mobile number already exist in database
    c.execute("SELECT * FROM users WHERE username=?", (username,))
    user = c.fetchone()
    if user:
        st.sidebar.warning('Username already exists. Please choose a different one.')
    else:
        c.execute("SELECT * FROM users WHERE mobile_number=?", (mobile_number,))
        user = c.fetchone()
        if user:
            st.sidebar.warning('Mobile number already exists. Please enter a different one.')
        else:
            if st.sidebar.button('Signup'):
                # Check if password matches confirmation
                if password == password_confirm:
                    # Add user to database
                    password_hash = hash_password(password)
                    c.execute("INSERT INTO users VALUES (?, ?, ?)", (username, password_hash, mobile_number))
                    conn.commit()
                    st.sidebar.success('Account created successfully!')
                else:
                    st.sidebar.error('Passwords do not match. Please try again.')


#------------------------------------------RESET PASSWORD FORM----------------------------------------------            
elif page == 'Reset Password/Username':
    mobile_number = st.sidebar.text_input('Mobile Number')
    if mobile_number:
        user = check_mobile_number(mobile_number)
        if user:
            st.sidebar.success('Mobile number verified')
            reset_option = st.sidebar.selectbox('Select what you want to reset', ['Password', 'Username'])
            if reset_option == 'Password':
                new_password = st.sidebar.text_input('New Password', type='password')
                confirm_password = st.sidebar.text_input('Confirm Password', type='password')
                if new_password and confirm_password:
                    if new_password == confirm_password:
                        # Update password hash in database
                        password_hash = hash_password(new_password)
                        c.execute("UPDATE users SET password_hash=? WHERE username=?", (password_hash, user))
                        conn.commit()
                        st.sidebar.success('Password updated successfully')
                    else:
                        st.sidebar.error('Passwords do not match')
            elif reset_option == 'Username':
                new_username = st.sidebar.text_input('New Username')
                if new_username:
                    # Check if username already exists in database
                    c.execute("SELECT * FROM users WHERE username=?", (new_username,))
                    existing_user = c.fetchone()
                    if existing_user:
                        st.sidebar.error('Username already exists. Please choose a different one.')
                    else:
                        c.execute("UPDATE users SET username=? WHERE mobile_number=?", (new_username, mobile_number))
                        conn.commit()
                        st.sidebar.success('Username updated successfully')
        else:
            st.sidebar.error('Mobile number not found')

#--------------------------------------------------------------------------------------------------------------
# Show app if user is logged in
if 'username' in st.experimental_get_query_params():
    username = st.experimental_get_query_params()['username'][0]
    st.subheader(f'Hi, {username} :wave:')
    


    

    # AssemblyAI API endpoint URLs
    upload_url = "https://api.assemblyai.com/v2/upload"
    transcript_url = "https://api.assemblyai.com/v2/transcript"

    # AssemblyAI API key
    API_KEY = "YOUR API KEY" #Enter the api key that you get from AssemblyAI 
   
    tab1, tab2 , tab3= st.tabs(["AUDIO TO TEXT AND VIDEO TO TEXT", "MY TRANSCRIPTIONS","ABOUT US"])

#--------------------------------------AUDIO/VIDEO TRANSCRIPTION SECTION-------------------------------------------------------- 
  
    # Initialize the languages that use to transcribe the audio/video to text
    with tab1:
        language = st.selectbox("Select language", ["English","Hindi","Spanish","French","German"])
        language_map = {        
            "English": "en",
            "Hindi": "hi",
            "Spanish": "es",
            "French": "fr",
            "German": "de" 
        }      

        # Container to show what files supported by audio and video 
        sup_container = st.container()

        col1, col2 = sup_container.columns(2)

        col1.write('Audio Supported Files : MP3, WAV, M4A')
        col2.write('Video Supported Files : MP4, MOV')

        # Upload file widget
        file = st.file_uploader("Choose an audio/video file", type=["mp3", "mp4", "m4a","wav", "mov"])

        # Check if a file was uploaded
        if file is not None:
            # Preview the audio file
            audio_bytes = file.read()
            st.audio(audio_bytes, format='audio/wav')

            #Save the filename to the database
            filename = file.name
            

            # Upload file to AssemblyAI API
            headers = {
                "authorization": API_KEY,
                "content-type": "application/json"
            }
            bar = st.progress(0)

            # Sending a request to upload URL
            upload_response = requests.post(
                upload_url,
                headers=headers,
                data=audio_bytes
            )

            # Extract the Upload URL from the response JSON
            audio_url = upload_response.json()["upload_url"]
            st.info('File has been uploaded to Textin')
            bar.progress(20)

            # Create transcription request
            json = {
                "audio_url": audio_url,                      
                "language_code" : language_map[language],
                "content_safety" : True             
            
            }

            # Submit transcription request to AssemblyAI
            response = requests.post(
                transcript_url,
                headers=headers,
                json=json
            )

            st.info('Transcribing uploaded file')
            bar.progress(40)

            # Extract transcript ID
            transcript_id = response.json()["id"]
            
            bar.progress(50)

            # Retrieve transcription results
            endpoint = f"https://api.assemblyai.com/v2/transcript/{transcript_id}"
            headers = {
                "authorization": API_KEY,
            }
            transcript_output_response = requests.get(endpoint, headers=headers)
            
            bar.progress(60)

            # Check if transcription is complete
            st.warning('Transcription is processing ...')
            while transcript_output_response.json()['status'] != 'completed':
                sleep(1)
                transcript_output_response = requests.get(endpoint, headers=headers)
                
            bar.progress(100)
            # Extract transcription text 
            transcription_text = transcript_output_response.json()["text"]          
          

            # Print transcribed text
            st.header('Output')

            with st.expander('Show Text'):
                st.success(transcription_text)

        
             # Write JSON to app
            with st.expander('Show Full Results'):
                st.write(transcript_output_response.json())
            
            # Write content_safety_labels
            with st.expander('Show content_safety_labels'):
                st.write(transcript_output_response.json()["content_safety_labels"])
            
            with st.expander('Summary of content_safety_labels'):
                st.write(transcript_output_response.json()["content_safety_labels"]["summary"])
                
        
            with st.expander('Download'):
                edited_text = st.text_area("Edit Transcription", value=transcription_text)
                if edited_text:
                    
                    
                    # Add download buttons
                    download_formats = ['txt','docx']
                    download_button_labels = ['Download as TXT','Download as Word']
                    download_button_keys = ['1','2']

                    for format, label in zip(download_formats, download_button_labels):
                        buffer = BytesIO()

                        if format == 'txt':
                            buffer.write(bytes(edited_text, 'utf-8'))
                            mime_type = 'text/plain'
                            file_extension = 'txt'

                        

                        elif format == 'docx':
                            document = Document()
                            document.add_heading('Transcription Text')                           
                            document.add_paragraph(edited_text)
                            document.save(buffer)
                            buffer.seek(0)
                            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                            file_extension = 'docx'

                        st.download_button(
                            label=label,
                            data=buffer,
                            file_name=f'transcription_text.{file_extension}',
                            mime=mime_type,
                        )
                
                    
                    
                    
             # Store the transcription data in the database
            c.execute("INSERT INTO transcripts (username,filename, language, text, transcript_status) VALUES (?, ?, ?, ?,  ?)", (username,filename, language, transcription_text, "completed"))
            conn.commit()

          

            st.success("Transcription completed and saved to database!")


                

#-------------------------------------------------MY TRANSCRIPTION SECTION------------------------------------------------
    with tab2:
       
        # Retrieve transcripts for the current user
        c.execute("SELECT * FROM transcripts WHERE username=?", (username,))
        transcripts = c.fetchall()


        
     
        # define functions for deleting and updating rows
        def delete_transcript(transcript_id):
            c.execute("DELETE FROM transcripts WHERE id=?", (transcript_id,))
            conn.commit()

        def update_transcript(transcript_id, new_data):
            c.execute("UPDATE transcripts SET filename=?, language=?, text=? WHERE id=?", (*new_data, transcript_id))
            conn.commit()

        # get the username of the currently logged-in user
        current_user = username

        # retrieve transcripts for the current user
        c.execute("SELECT id, filename, language, text FROM transcripts WHERE username=?", (current_user,))
        transcripts = c.fetchall()

        if transcripts:
            st.header(f"{current_user}'s Transcriptions")
            transcript_ids = [transcript[1] for transcript in transcripts]
            selected_transcript_id = st.selectbox("Select Transcript File", transcript_ids)
            selected_transcript = next((t for t in transcripts if t[1] == selected_transcript_id), None)
            if selected_transcript:
                with st.expander("View Language"):
                    st.write(selected_transcript[2])
                with st.expander("View Text"):
                    st.write(selected_transcript[3])
                with st.expander("Update Transcription"):                    
                    new_filename = st.text_input("Filename", selected_transcript[1])
                    new_language = st.text_input("Language", selected_transcript[2])
                    new_text = st.text_area("Text", selected_transcript[3])                    
                    if st.button("Update"):
                        try:
                            update_transcript(selected_transcript[0], (new_filename, new_language, new_text))
                            st.success("Transcript updated successfully!")
                        except:
                            st.error("Error updating transcript")                 

                with st.expander("Delete Transcription"):
                    if st.button("Delete"):
                        try:
                            delete_transcript(selected_transcript[0])
                            st.success("Transcript deleted successfully!")
                        except:
                            st.error("Error deleting transcript")

                with st.expander('Download'):
                
                    
                    # Add download buttons
                    download_formats = ['txt', 'docx']
                    download_button_labels = ['Download as TXT', 'Download as Word']
                    download_button_keys = ['3','4']

                    for format, label,key in zip(download_formats, download_button_labels, download_button_keys):
                        buffer = BytesIO()

                        if format == 'txt':
                            buffer.write(bytes(new_text, 'utf-8'))
                            mime_type = 'text/plain'
                            file_extension = 'txt'

                        

                        elif format == 'docx':
                            document = Document()
                            document.add_heading('Transcription Text')
                            document.add_paragraph(new_text)
                            document.save(buffer)
                            buffer.seek(0)
                            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                            file_extension = 'docx'

          
                        st.download_button(
                            label=label,
                            data=buffer,
                            file_name=f'transcription_text.{file_extension}',
                            mime=mime_type,
                            key=key,
                           
                        )
 


            else:
                st.warning("Transcript not found")
        else:
            st.warning("No transcripts found for this user")

        
#-----------------------------------------ABOUTUS TAB SECTION-------------------------------------------------------------
    with tab3:        
        with st.container():
            st.write("---")
            left_column, right_columns = st.columns(2)
            with left_column:
                st.subheader("What Textin Can Do for You :question:")
                st.write("##")
                st.write(
                    """ 
                    Textin is a Transcription Website that uses Speech Recognition technology to convert voice                  
                    into text.
                    - It provides flexibility to the user regarding transcription.
                    - By using Textin users can easily convert their records as per their needs.
                    - It also supports different audio/video file formats such as mp3,mp4, and wav, 
                      etc for conversion.
                    - It also supports different languages such as English, Hindi, etc
                    - It also has a feature in which users can save and download the converted text into txt and word. 
                    - After the conversion of the audio/video  file, the user can also edit the converted text.
                    - It also provides the storage of the converted text file in the database.
                    - By using Textin, users will able to save time efficiently.

                    """
                )
                with right_columns:
                    image = st.image('images/about.png', width=None)



    # close the connection
        conn.commit()
        conn.close()
    



# Logout button
if st.sidebar.button('Logout'):
    # Clear username from query parameters
    app_state = st.experimental_get_query_params()
    app_state.pop('username', None)
    st.experimental_set_query_params(**app_state)
    # Show login page
    st.experimental_rerun()



